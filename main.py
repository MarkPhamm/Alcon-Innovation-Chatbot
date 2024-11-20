import streamlit as st
import requests
from bs4 import BeautifulSoup
from parse import parse_with_gpt4_stream  # Import the correct parsing function
import io  # To create file-like objects
from docx import Document  # For creating Word documents
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# API configuration
API_KEY = os.getenv('API_KEY')
SCRAPER_API_URL = os.getenv('SCRAPER_API_URL')

# Predefined URLs
URL_OPTIONS = {
    "JnJ": "https://www.jnj.com/media-center/press-releases/johnson-johnson-rolls-out-new-tecnis-odyssey-next-generation-intraocular-lens-offering-cataract-patients-precise-vision-at-every-distance-in-any-lighting",
    # "RXSight": "https://www.centerforeyes.com/rxsight/",  # Example link
    "Zeiss IOLs": "https://www.zeiss.com/meditec-ag/en/media-news/press-releases/2024/zeiss-at-aao.html"  # Example link
}

# Function to clean HTML content
def clean_body_content(body_content):
    soup = BeautifulSoup(body_content, "html.parser")

    for script_or_style in soup(["script", "style"]):
        script_or_style.extract()

    # Get text or further process the content
    cleaned_content = soup.get_text(separator="\n")
    cleaned_content = "\n".join(
        line.strip() for line in cleaned_content.splitlines() if line.strip()
    )

    return cleaned_content

# Function to scrape website
def scrape_website(url):
    payload = {'api_key': API_KEY, 'url': url}
    response = requests.get(SCRAPER_API_URL, params=payload)
    
    if response.status_code == 200:
        return response.text
    else:
        st.error(f"Failed to scrape the website. Status code: {response.status_code}")
        return ""

# Function to split content into chunks (assuming a chunk size for OpenAI GPT input)
def split_content(content, chunk_size=5000):
    return [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]

# Streamlit UI
st.title("Innov8 AI Platform")

# Multi-select dropdown to choose websites
selected_sites = st.multiselect("Select Websites to Scrape", options=list(URL_OPTIONS.keys()))

# Step 1: Scrape the Website(s)
if st.button("Scrape Selected Websites"):
    if selected_sites:
        st.write("Scraping the selected websites...")
        
        scraped_content = ""
        # Loop through selected sites and scrape each
        for site in selected_sites:
            url = URL_OPTIONS[site]
            site_content = scrape_website(url)
            if site_content:
                cleaned_content = clean_body_content(site_content)
                scraped_content += f"\n\n--- Content from {site} ---\n\n" + cleaned_content
        
        # Store the cleaned content in Streamlit session state
        if scraped_content:
            st.session_state.scraped_content = scraped_content
            
            # Display the cleaned content in an expandable text box
            with st.expander("View Cleaned Content"):
                st.text_area("Cleaned Content", scraped_content, height=300)

if "scraped_content" in st.session_state:
    parse_description = st.text_area("Describe what you want to parse")

    if st.button("Parse Content"):
        if parse_description:
            st.write("Parsing the content...")

            # Split the content into chunks for OpenAI GPT processing
            dom_chunks = split_content(st.session_state.scraped_content)
            
            # Parse the content with OpenAI GPT
            parsed_result = parse_with_gpt4_stream(dom_chunks, parse_description)
            
            if parsed_result.strip():  # Check if the result is not empty
                st.write(parsed_result)
            else:
                st.write("No matching information found based on the provided description.")

# Add a download button for the parsed result
if "parsed_result" in st.session_state:
    # Creating a file-like object for the parsed content in plain text
    parsed_content_io = io.StringIO(st.session_state.parsed_result)

    # Creating a download button for plain text file
    st.download_button(
        label="Download Parsed Content (Text File)",
        data=parsed_content_io.getvalue(),  # Get the content as a string
        file_name="parsed_content.txt",
        mime="text/plain"
    )

    # Creating a Word document with the parsed content
    doc = Document()
    doc.add_heading('Parsed Content', 0)
    doc.add_paragraph(st.session_state.parsed_result)
    
    # Save the Word document to an in-memory file
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Reset the pointer to the beginning of the file

    # Creating a download button for the Word document
    st.download_button(
        label="Download Parsed Content (Word Document)",
        data=doc_io,  # File-like object for the Word document
        file_name="parsed_content.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )